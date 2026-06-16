#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Builds a Word document summarising SAR scenes from:
 A) MERIA: D:\Masters\MERIA\raw_grd\**\*.SAFE\measurement\*.tif*
 B) MARIDA: D:\Masters\MARIDA\downloads\**\SAR_*h\*.tif

Outputs:
 - Two PNG maps with centroids
 - A .docx with:
    Section A (MERIA): map + table
    Section B (MARIDA): map + table
    Section C (Planet API cURLs for nearest PSScene within ±12 h)

Author: ChatGPT
"""
import os
import re
import sys
import glob
import math
import json
import traceback
from datetime import datetime, timedelta, timezone
from pathlib import Path

# --- Third-party deps (install if missing) ---
try:
    import rasterio
    from rasterio.warp import transform_bounds
except Exception as e:
    print("ERROR: rasterio is required. Install with: pip install rasterio", file=sys.stderr)
    raise

try:
    from pyproj import CRS, Transformer
except Exception as e:
    print("ERROR: pyproj is required. Install with: pip install pyproj", file=sys.stderr)
    raise

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
except Exception as e:
    print("ERROR: matplotlib is required. Install with: pip install matplotlib", file=sys.stderr)
    raise

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except Exception as e:
    print("ERROR: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    raise

# ---------- CONFIG ----------
MERIA_ROOT = r"D:\Masters\MERIA\raw_grd"
MARIDA_ROOT = r"D:\Masters\MARIDA\downloads"
OUTPUT_DOCX = r"D:\Masters\SAR_inventory_planet.docx"
OUTDIR = str(Path(OUTPUT_DOCX).parent)
MAP_MERIA_PNG = os.path.join(OUTDIR, "map_meria_points.png")
MAP_MARIDA_PNG = os.path.join(OUTDIR, "map_marida_points.png")

# Planet search window around SAR time
PLANET_WINDOW_HOURS = 12  # ±12 h

# ---------- HELPERS ----------
def parse_datetime_lower(s: str):
    """
    Accepts strings like 20210730t163253 (lowercase 't') and returns aware UTC datetime.
    """
    s = s.strip()
    try:
        return datetime.strptime(s, "%Y%m%dt%H%M%S").replace(tzinfo=timezone.utc)
    except:
        return None

def parse_datetime_upper(s: str):
    """
    Accepts strings like 20210730T163253 (uppercase 'T') and returns aware UTC datetime.
    """
    s = s.strip()
    try:
        return datetime.strptime(s, "%Y%m%dT%H%M%S").replace(tzinfo=timezone.utc)
    except:
        return None

def safe_float(v, ndp=6):
    try:
        return round(float(v), ndp)
    except:
        return None

def centroid_lonlat_from_raster(path: str):
    """
    Returns (lon, lat) centroid in WGS84 from raster bounds/CRS.
    """
    with rasterio.open(path) as ds:
        b = ds.bounds
        if ds.crs is None:
            # Assume WGS84 if missing
            lon = (b.left + b.right) / 2.0
            lat = (b.top + b.bottom) / 2.0
            return float(lon), float(lat)
        crs = CRS.from_user_input(ds.crs)
        # bounds may be in projected CRS; transform to WGS84
        try:
            tb = transform_bounds(crs, "EPSG:4326", b.left, b.bottom, b.right, b.top, densify_pts=21)
            lon = (tb[0] + tb[2]) / 2.0
            lat = (tb[1] + tb[3]) / 2.0
            return float(lon), float(lat)
        except Exception:
            # Fallback via pixel center
            cx = ds.transform * (ds.width/2.0, ds.height/2.0)
            # transform pixel center to WGS84
            try:
                transformer = Transformer.from_crs(crs, CRS.from_epsg(4326), always_xy=True)
                lon, lat = transformer.transform(cx[0], cx[1])
                return float(lon), float(lat)
            except Exception:
                raise

def scan_meria(root: str):
    """
    MERIA: find measurement tiffs under *.SAFE/measurement/*.tif*
    Extract:
      - product (SAFE folder name)
      - polarisation (vv/vh if in filename)
      - start/end time (from filename or SAFE folder)
      - centroid lon/lat
    """
    pattern1 = os.path.join(root, "**", "*.SAFE", "measurement", "*.tif*")
    files = glob.glob(pattern1, recursive=True)
    out = []
    for f in files:
        try:
            p = Path(f)
            # SAFE folder name
            safe_dir = None
            for parent in p.parents:
                if parent.name.endswith(".SAFE"):
                    safe_dir = parent.name
                    break

            # pol from filename e.g. ...-vh-...tif
            pol = None
            mpol = re.search(r'-(vv|vh)-', p.name, re.IGNORECASE)
            if mpol:
                pol = mpol.group(1).upper()

            # times: prefer from file name: s1b-iw-grd-vh-20210730t163253-20210730t163318-...
            start_dt = end_dt = None
            mfile = re.search(r'-(\d{8}t\d{6})-(\d{8}t\d{6})-', p.name, re.IGNORECASE)
            if mfile:
                start_dt = parse_datetime_lower(mfile.group(1))
                end_dt   = parse_datetime_lower(mfile.group(2))

            # fallback from SAFE folder: ..._20210730T163253_20210730T163318_...
            if start_dt is None or end_dt is None:
                if safe_dir:
                    msafe = re.search(r'_(\d{8}T\d{6})_(\d{8}T\d{6})_', safe_dir)
                    if msafe:
                        start_dt = start_dt or parse_datetime_upper(msafe.group(1))
                        end_dt   = end_dt   or parse_datetime_upper(msafe.group(2))

            lon, lat = centroid_lonlat_from_raster(f)
            out.append({
                "source": "MERIA",
                "path": f,
                "product": safe_dir or p.stem,
                "file": p.name,
                "pol": pol or "",
                "start": start_dt,
                "end": end_dt,
                "lon": safe_float(lon),
                "lat": safe_float(lat),
            })
        except Exception as e:
            print(f"[MERIA] Skipping {f} due to error: {e}", file=sys.stderr)
    return out

def scan_marida(root: str):
    """
    MARIDA: inside downloads/**/SAR_* (plus/minus hours) folders; get *.tif
    Filenames like: S1_16PDC_20181024T113758_vh.tif
    Extract:
      - product (folder name one level above the file or stem)
      - polarisation (_vh or _vv)
      - acquired time (from filename)
      - centroid lon/lat
    """
    # Find all tif under folders named SAR_* (e.g., SAR_-4.6h)
    pattern = os.path.join(root, "**", "SAR_*", "*.tif")
    files = glob.glob(pattern, recursive=True)
    out = []
    for f in files:
        try:
            p = Path(f)
            pol = ""
            mpol = re.search(r'_(vv|vh)\.tif[f]?$', p.name, re.IGNORECASE)
            if mpol:
                pol = mpol.group(1).upper()
            # time from filename: ..._20181024T113758_...
            mtime = re.search(r'_(\d{8}T\d{6})_', p.name)
            acquired = parse_datetime_upper(mtime.group(1)) if mtime else None

            # product label: use folder containing the SAR_* dir and date
            # e.g., ...\downloads\16PDC\2018-10-24\SAR_-4.6h\file.tif
            # Let's call product = join of the two parents above SAR_* (e.g., "16PDC_2018-10-24")
            parent = p.parent  # SAR_* dir
            prod_label = ""
            try:
                if parent and parent.parent:
                    prod_label = f"{parent.parent.name}_{parent.name}"
            except:
                prod_label = p.stem

            lon, lat = centroid_lonlat_from_raster(f)
            out.append({
                "source": "MARIDA",
                "path": f,
                "product": prod_label or p.stem,
                "file": p.name,
                "pol": pol or "",
                "start": acquired,   # use 'start' field for single timestamp
                "end": None,
                "lon": safe_float(lon),
                "lat": safe_float(lat),
            })
        except Exception as e:
            print(f"[MARIDA] Skipping {f} due to error: {e}", file=sys.stderr)
    return out

def make_points_map(points, out_png, title):
    """
    Make a simple lon/lat scatter plot. No internet or base tiles required.
    points: list of dicts with keys 'lon','lat'
    """
    if not points:
        # create a placeholder figure
        fig = plt.figure(figsize=(8, 4.5), dpi=150)
        ax = plt.gca()
        ax.set_title(f"{title} (no points found)")
        ax.set_xlabel("Longitude")
        ax.set_ylabel("Latitude")
        fig.tight_layout()
        fig.savefig(out_png)
        plt.close(fig)
        return

    lons = [d["lon"] for d in points if d.get("lon") is not None]
    lats = [d["lat"] for d in points if d.get("lat") is not None]

    fig = plt.figure(figsize=(8, 4.5), dpi=150)
    ax = plt.gca()
    ax.scatter(lons, lats, s=12)
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    ax.set_title(title)

    # set a reasonable extent with padding
    min_lon, max_lon = min(lons), max(lons)
    min_lat, max_lat = min(lats), max(lats)
    pad_lon = max(0.2, (max_lon - min_lon) * 0.1) if max_lon > min_lon else 1.0
    pad_lat = max(0.2, (max_lat - min_lat) * 0.1) if max_lat > min_lat else 1.0
    ax.set_xlim(min_lon - pad_lon, max_lon + pad_lon)
    ax.set_ylim(min_lat - pad_lat, max_lat + pad_lat)

    # crude aspect correction for lon/lat scaling
    mean_lat = (min_lat + max_lat) / 2.0
    try:
        ax.set_aspect(1.0 / math.cos(math.radians(mean_lat)))
    except Exception:
        pass

    fig.tight_layout()
    fig.savefig(out_png)
    plt.close(fig)

def dt_to_str(dt):
    if dt is None:
        return ""
    # Always output as UTC ISO8601 Z
    return dt.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

def planet_curl_for_point_psscene(lon, lat, center_dt, window_hours=12):
    """
    Builds a cURL for Planet Data API v1 quick-search, limiting to PSScene items
    within ±window_hours around center_dt at a point geometry. Returns a string.
    """
    if center_dt is None:
        # if no timestamp, give a same-day wide window as a fallback
        gte = "1970-01-01T00:00:00Z"
        lte = "2100-01-01T00:00:00Z"
    else:
        start = (center_dt - timedelta(hours=window_hours)).astimezone(timezone.utc)
        end   = (center_dt + timedelta(hours=window_hours)).astimezone(timezone.utc)
        gte = start.strftime("%Y-%m-%dT%H:%M:%SZ")
        lte = end.strftime("%Y-%m-%dT%H:%M:%SZ")

    payload = {
        "item_types": ["PSScene"],
        "filter": {
            "type": "AndFilter",
            "config": [
                {
                    "type": "GeometryFilter",
                    "field_name": "geometry",
                    "config": {"type": "Point", "coordinates": [float(lon), float(lat)]}
                },
                {
                    "type": "DateRangeFilter",
                    "field_name": "acquired",
                    "config": {"gte": gte, "lte": lte}
                }
            ]
        },
        "sort": [{"field_name": "acquired", "direction": "asc"}],
        "limit": 1
    }
    # Note: user must set PL_API_KEY env var or replace inline
    curl = (
        'curl -u "$PL_API_KEY:" -X POST "https://api.planet.com/data/v1/quick-search" '
        '-H "Content-Type: application/json" '
        f"-d '{json.dumps(payload, separators=(',', ':'))}'"
    )
    return curl

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_table(doc, headers, rows, col_widths_in=None, monospace_cols=None):
    """
    headers: list[str]
    rows: list[list[str]]
    col_widths_in: list[float] | None  (inches)
    monospace_cols: set[int] | None
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for r in rows:
        row_cells = table.add_row().cells
        for j, val in enumerate(r):
            para = row_cells[j].paragraphs[0]
            run = para.add_run(str(val))
            if monospace_cols and j in monospace_cols:
                run.font.name = "Courier New"
                r_elm = run._element
                r_pr = r_elm.rPr
                r_fonts = r_pr.rFonts
                r_fonts.set(qn('w:eastAsia'), 'Courier New')
                run.font.size = Pt(9)
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            try:
                table.columns[i].width = Inches(w)
            except Exception:
                pass
    doc.add_paragraph()  # spacing
    return table

def build_doc(meria_list, marida_list, map_meria_png, map_marida_png, out_docx):
    doc = Document()
    doc.add_heading("SAR Image Inventory – MERIA & MARIDA", level=0)
    p = doc.add_paragraph("Generated by script. Times are UTC. Coordinates are centroid (lon, lat).")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # SECTION A: MERIA
    add_heading(doc, "Section A: MERIA", level=1)
    doc.add_paragraph("Locations of MERIA SAR images:")
    if os.path.isfile(map_meria_png):
        doc.add_picture(map_meria_png, width=Inches(6.5))
    else:
        doc.add_paragraph("(map not available)")
    doc.add_paragraph()

    headers_meria = ["#", "Product (SAFE)", "File", "Pol", "Start (UTC)", "End (UTC)", "Centroid (lon, lat)"]
    rows_meria = []
    for i, d in enumerate(sorted(meria_list, key=lambda x: (x.get("start") or datetime.min.replace(tzinfo=timezone.utc)))):
        rows_meria.append([
            i + 1,
            d.get("product", ""),
            d.get("file", ""),
            d.get("pol", ""),
            dt_to_str(d.get("start")),
            dt_to_str(d.get("end")),
            f'{d.get("lon")}, {d.get("lat")}'
        ])
    add_table(doc, headers_meria, rows_meria, col_widths_in=[0.4, 2.4, 2.0, 0.5, 1.2, 1.2, 1.4])

    # SECTION B: MARIDA
    add_heading(doc, "Section B: MARIDA", level=1)
    doc.add_paragraph("Locations of MARIDA SAR images:")
    if os.path.isfile(map_marida_png):
        doc.add_picture(map_marida_png, width=Inches(6.5))
    else:
        doc.add_paragraph("(map not available)")
    doc.add_paragraph()

    headers_marida = ["#", "Product", "File", "Pol", "Acquired (UTC)", "Centroid (lon, lat)"]
    rows_marida = []
    for i, d in enumerate(sorted(marida_list, key=lambda x: (x.get("start") or datetime.min.replace(tzinfo=timezone.utc)))):
        rows_marida.append([
            i + 1,
            d.get("product", ""),
            d.get("file", ""),
            d.get("pol", ""),
            dt_to_str(d.get("start")),
            f'{d.get("lon")}, {d.get("lat")}'
        ])
    add_table(doc, headers_marida, rows_marida, col_widths_in=[0.4, 2.4, 2.4, 0.6, 1.3, 1.4])

    # SECTION C: Planet API Calls
    add_heading(doc, "Section C: Planet API (Nearest PSScene within ±12 h)", level=1)
    doc.add_paragraph(
        'These cURL commands query Planet Data API v1 "quick-search" for PSScene nearest to the SAR time. '
        'Set environment variable PL_API_KEY to your Planet API key before running.'
    )

    headers_api = ["#", "Source", "Product", "Time (UTC)", "Centroid (lon,lat)", "Quick Search cURL"]
    rows_api = []
    all_items = []
    for d in meria_list:
        all_items.append(("MERIA", d))
    for d in marida_list:
        all_items.append(("MARIDA", d))

    for i, (src, d) in enumerate(sorted(all_items, key=lambda x: (x[1].get("start") or datetime.min.replace(tzinfo=timezone.utc)))):
        lon, lat = d.get("lon"), d.get("lat")
        curl = planet_curl_for_point_psscene(lon, lat, d.get("start"), PLANET_WINDOW_HOURS)
        rows_api.append([
            i + 1,
            src,
            d.get("product", ""),
            dt_to_str(d.get("start")),
            f"{lon}, {lat}",
            curl
        ])
    # Make last column monospace for readability
    add_table(doc, headers_api, rows_api,
              col_widths_in=[0.4, 0.9, 1.8, 1.4, 1.3, 3.5],
              monospace_cols={5})

    doc.save(out_docx)

def main():
    # Allow optional CLI overrides:
    # python script.py [MERIA_ROOT] [MARIDA_ROOT] [OUTPUT_DOCX]
    meria = MERIA_ROOT
    marida = MARIDA_ROOT
    outdoc = OUTPUT_DOCX
    if len(sys.argv) >= 2:
        meria = sys.argv[1]
    if len(sys.argv) >= 3:
        marida = sys.argv[2]
    if len(sys.argv) >= 4:
        outdoc = sys.argv[3]

    os.makedirs(os.path.dirname(outdoc), exist_ok=True)

    print("Scanning MERIA ...")
    meria_list = scan_meria(meria)
    print(f"  MERIA scenes: {len(meria_list)}")

    print("Scanning MARIDA ...")
    marida_list = scan_marida(marida)
    print(f"  MARIDA scenes: {len(marida_list)}")

    # Build maps
    print("Making maps ...")
    make_points_map(meria_list, MAP_MERIA_PNG, "MERIA SAR Scene Centroids")
    make_points_map(marida_list, MAP_MARIDA_PNG, "MARIDA SAR Scene Centroids")

    # Build docx
    print(f"Writing Word document: {outdoc}")
    build_doc(meria_list, marida_list, MAP_MERIA_PNG, MAP_MARIDA_PNG, outdoc)

    print("Done.")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print("FATAL:", e, file=sys.stderr)
        traceback.print_exc()
        sys.exit(1)
