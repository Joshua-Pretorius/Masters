from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\Masters\Mock_Literature_Review_Section.docx")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_callout(doc, text):
    p = doc.add_paragraph(style="Mock Note")
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(14)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(text)
    r.italic = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text, style="Normal")
    p.paragraph_format.widow_control = True
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Resolved preset: narrative_proposal.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for name, size, color, before, after in (
    ("Heading 1", 16, "2E74B5", 18, 10),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

if "Mock Note" not in [s.name for s in styles]:
    note_style = styles.add_style("Mock Note", WD_STYLE_TYPE.PARAGRAPH)
else:
    note_style = styles["Mock Note"]
note_style.font.name = "Calibri"
note_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
note_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
note_style.font.size = Pt(10)
note_style.font.color.rgb = RGBColor(60, 70, 82)
note_style.paragraph_format.left_indent = Inches(0.18)
note_style.paragraph_format.right_indent = Inches(0.18)
note_style.paragraph_format.space_before = Pt(6)
note_style.paragraph_format.space_after = Pt(6)
note_style.paragraph_format.line_spacing = 1.15

# Quiet running furniture for a multi-page academic mock-up.
header = section.header.paragraphs[0]
header.text = "ILLUSTRATIVE LITERATURE REVIEW EXCERPT"
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_run = header.runs[0]
header_run.font.name = "Calibri"
header_run.font.size = Pt(8.5)
header_run.font.bold = True
header_run.font.color.rgb = RGBColor(115, 115, 115)
add_page_number(section.footer.paragraphs[0])

# Compact editorial-cover pattern, intentionally adapted for a short working excerpt.
kicker = doc.add_paragraph()
kicker.paragraph_format.space_before = Pt(8)
kicker.paragraph_format.space_after = Pt(6)
kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = kicker.add_run("WORKING MOCK-UP")
r.font.name = "Calibri"
r.font.size = Pt(9)
r.font.bold = True
r.font.color.rgb = RGBColor(46, 116, 181)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(6)
r = title.add_run("From the Data Bottleneck to an Appropriate Learning Strategy")
r.font.name = "Calibri"
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor(31, 77, 120)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after = Pt(18)
r = subtitle.add_run("Illustrative Sections 2.5-2.8 for a SAR marine-debris literature review")
r.font.name = "Calibri"
r.font.size = Pt(12.5)
r.font.italic = True
r.font.color.rgb = RGBColor(85, 85, 85)

add_callout(
    doc,
    "Purpose of this mock-up: to demonstrate how the literature can move from available data, "
    "to label creation, and only then to machine-learning strategies suited to scarce and uncertain annotations. "
    "This is a separate illustrative excerpt and has not been inserted into the existing thesis draft."
)

add_heading(doc, "2.5 Labelled Data for SAR-Based Marine-Debris Detection", 1)

add_body(doc, "Machine-learning methods for floating marine-debris detection depend on labelled observations that connect a remotely sensed signal to a known target. In optical imagery, this connection can sometimes be supported by spectral contrast, visual interpretation or observations of controlled targets. In SAR imagery, however, the measured response is indirect. Floating debris may alter local surface roughness or occur within convergence features, but similar responses can also be produced by wind variability, fronts, oil slicks, biological material and other floating objects. A useful SAR dataset must therefore do more than identify anomalous pixels: it must represent how confidently those anomalies can be attributed to marine debris and under which environmental conditions they were observed.")

add_body(doc, "This distinction is important because the target of detection is not necessarily an individual plastic item. Medium-resolution satellite imagery is more likely to represent spatially coherent accumulations of floating macroplastic and associated marine debris, such as patches or windrows, than dispersed particles. The required labels must consequently match the spatial unit that a model is expected to predict. A dataset designed for patch classification, for example, may assign one label to an image window, whereas a segmentation dataset requires boundaries that indicate where the accumulation is believed to occur. The reliability of those boundaries directly affects what can defensibly be learned and evaluated.")

add_heading(doc, "2.5.1 Existing Labelled Datasets", 2)

add_body(doc, "Existing marine-debris datasets are dominated by optical observations. MARIDA provides multiclass Sentinel-2 annotations that distinguish marine debris from several co-occurring ocean-surface features and has enabled the development of optical segmentation baselines (Kikaki et al., 2022). Controlled experiments undertaken through the Plastic Litter Project have also demonstrated the value of placing artificial targets at sea during satellite acquisitions, thereby creating observations for which target presence is known more reliably (Topouzelis et al., 2020). These resources establish useful precedents for label design, particularly the inclusion of confusing background classes and the use of independent observations to support interpretation.")

add_body(doc, "Their direct applicability to SAR remains limited. Optical sensors respond to reflected electromagnetic energy and can exploit spectral differences between water and floating materials, while SAR primarily responds to surface geometry, roughness and dielectric properties. An optical marine-debris label therefore does not guarantee that the same feature will be visible, separable or spatially coincident in a SAR acquisition. Differences in acquisition time introduce an additional problem because floating material may move between observations. Optical datasets can support the creation of candidate SAR labels, but they should not be treated as interchangeable SAR ground truth.")

add_heading(doc, "2.5.2 The SAR Data Bottleneck", 2)

add_body(doc, "Only a small number of studies have directly investigated labelled SAR data for plastic or mixed marine debris. Savastano et al. (2021) identified debris in Sentinel-2 imagery and used nearby Sentinel-1 acquisitions to create a preliminary SAR dataset for shallow classification. Other studies have examined related floating features or controlled targets to determine when debris produces observable SAR responses (Qi et al., 2022; de Fockert et al., 2024). Collectively, this work supports the feasibility of SAR-based investigation, but it also shows that the signal is conditional and that available labelled samples remain limited.")

add_body(doc, "The shortage is therefore not only a question of dataset size. It is also a shortage of labels that preserve information about temporal separation, environmental conditions, target visibility and uncertainty. Increasing the number of image patches without increasing the number and diversity of independent scenes would not necessarily resolve this limitation, because neighbouring patches may contain nearly identical acquisition and sea-state characteristics. The relevant evidence base is determined by the diversity of independently observed events, environments and sensor configurations, rather than by the number of derivative image crops alone.")

add_heading(doc, "2.6 Creating Labels for SAR Imagery", 1)

add_body(doc, "A SAR training dataset can be developed through direct observation, controlled deployments, manual interpretation or transfer from another sensor. Direct field observations and controlled targets provide the clearest evidence of target presence, but they are expensive and difficult to coordinate with satellite overpasses. Manual interpretation of SAR alone is also problematic because debris does not have a single, consistent radar signature. Optical-to-SAR transfer offers a practical way to increase coverage, provided that the transfer is treated as an uncertain inference rather than as a direct observation.")

add_heading(doc, "2.6.1 Optical-to-SAR Label Transfer", 2)

add_body(doc, "In an optical-to-SAR workflow, a candidate debris accumulation is first identified in an optical image using spectral indices, visual interpretation or an existing annotation. A corresponding SAR acquisition is then located within an acceptable temporal window. Where the two acquisitions are closely timed, the optical geometry may provide an initial estimate of the target position in the SAR image. This approach follows the broad precedent established by Savastano et al. (2021), while allowing a substantially larger set of paired observations to be assembled.")

add_body(doc, "The transfer cannot be based on spatial overlay alone. Floating debris is transported by surface currents, wind-induced drift and waves, and its geometry may disperse or reorganise between acquisitions. A label copied directly from an earlier optical image may therefore be spatially inaccurate even when both images are correctly georeferenced. The temporal gap, predicted displacement and uncertainty in that displacement should be retained as properties of the label. Drift modelling may improve the estimated position, but the predicted correction should itself be validated rather than assumed to be more accurate.")

add_heading(doc, "2.6.2 Representing Label Confidence", 2)

add_body(doc, "A binary distinction between debris and non-debris can conceal substantial differences in evidential quality. A more defensible dataset could distinguish high-confidence cores, probable debris, uncertain boundaries and regions that should be ignored during training. Confidence may depend on the clarity of the optical detection, the temporal gap between acquisitions, estimated drift uncertainty, agreement between independent annotators and whether a corresponding response is visible in the SAR image. This structure permits uncertain information to be retained without presenting every transferred pixel as equally reliable.")

add_body(doc, "Negative data require similar care. Random ocean pixels may create an unrealistically easy background class, while the operational challenge lies in separating debris from ships, wakes, slicks, fronts, whitecaps, current lines and biological aggregations. These hard-negative classes should be deliberately represented, together with environmental metadata such as wind and sea-state conditions. A smaller independently validated subset should then be reserved to test whether models trained on proxy labels generalise to observations with stronger evidence of target presence.")

add_heading(doc, "2.7 Machine Learning with Scarce, Imbalanced and Uncertain Labels", 1)

add_body(doc, "The properties of the resulting dataset should determine the learning strategy. The problem is not simply that few annotations are available; positive examples are rare, labels may be spatially uncertain, and numerous ocean features can resemble the proposed target. A model selected without accounting for these characteristics may learn acquisition-specific patterns or annotation artefacts rather than transferable evidence of marine debris.")

add_heading(doc, "2.7.1 Task Formulation and Baselines", 2)

add_body(doc, "Patch classification offers a comparatively simple formulation in which each image window is assigned a debris or non-debris label. This may be appropriate where transferred boundaries are unreliable, but it provides limited information about the spatial extent of detected accumulations. Semantic segmentation can preserve the geometry of patches and windrows, although it requires sufficiently credible pixel-level masks. Comparing a simple patch-based baseline with a segmentation model would help determine whether the additional spatial detail is supported by the available annotations rather than assuming that a more complex model is necessarily superior.")

add_heading(doc, "2.7.2 Learning from Limited and Imperfect Supervision", 2)

add_body(doc, "Deep-learning models generally require more labelled examples than are available in emerging SAR marine-debris datasets. Data augmentation can increase variation in orientation, position and intensity, but it does not create new independent environmental events. Transfer learning and self-supervised learning offer a more relevant response by allowing an encoder to learn general representations from larger volumes of unlabelled SAR imagery before being fine-tuned on the smaller labelled dataset. Domain-specific pretraining is particularly attractive because the model can first learn common ocean-surface structures without requiring debris annotations (Li et al., 2024).")

add_body(doc, "Training should also reflect class imbalance and label confidence. Loss functions that emphasise rare positive regions may reduce domination by the background class, while uncertain boundary pixels can be down-weighted or excluded. Weakly supervised and semi-supervised approaches may allow high-confidence labels to guide learning while lower-confidence or unlabelled observations provide additional context. The value of these methods should nevertheless be established experimentally, as additional complexity cannot compensate for systematically incorrect transferred labels.")

add_heading(doc, "2.7.3 Evaluation and Generalisation", 2)

add_body(doc, "Evaluation must separate independent scenes or debris events rather than randomly dividing neighbouring patches. Random patch splits risk placing near-duplicate samples from one acquisition in both training and testing data, producing optimistic estimates of generalisation. Where sufficient data are available, geographic or temporal hold-outs can provide stronger evidence that a model transfers beyond the conditions represented during training.")

add_body(doc, "Accuracy alone is unlikely to be informative for a strongly imbalanced problem because a model can achieve a high score by predicting the dominant water class. Precision, recall, F1 score and intersection-over-union provide more direct evidence of positive-class performance, while event-level assessment can indicate whether a spatial accumulation was detected at all. Results should also be stratified by label confidence, temporal gap and environmental conditions. Such analysis would show not only whether detection is possible, but also the circumstances under which model predictions cease to be reliable.")

add_heading(doc, "2.8 Critical Synthesis", 1)

add_body(doc, "The literature indicates that floating marine debris can produce observable responses in SAR imagery, but those responses are neither unique nor consistent across sea states. Progress is constrained by the limited availability of SAR annotations and by uncertainty introduced when labels are derived from optical observations acquired at different times. The principal research problem is therefore not merely the selection of a deep-learning architecture. It is the construction of a sufficiently transparent evidence chain connecting an optical observation, a potentially displaced target, its conditional SAR response and the model used to detect it.")

add_body(doc, "A defensible contribution would combine an uncertainty-aware, optically derived SAR dataset with an evaluation of learning strategies suited to scarce, imbalanced and imperfect labels. This framing permits the data-creation process and the machine-learning experiments to address the same research gap. It also places limits on the resulting claims: the study would assess the conditions under which floating marine-debris accumulations are detectable in SAR, rather than assuming that SAR provides direct or universally reliable identification of plastic.")

add_heading(doc, "References Cited in This Mock-Up", 1)
references = [
    "de Fockert, A. et al. (2024) 'Assessing the detection of floating plastic litter with advanced remote sensing technologies in a hydrodynamic test facility', Scientific Reports, 14, 25902.",
    "Kikaki, K. et al. (2022) 'MARIDA: A benchmark for Marine Debris detection from Sentinel-2 remote sensing data', PLOS ONE, 17(1), e0262247.",
    "Li, W. et al. (2024) 'Predicting gradient is better: Exploring self-supervised learning for SAR ATR with a joint-embedding predictive architecture', ISPRS Journal of Photogrammetry and Remote Sensing, 218, pp. 326-338.",
    "Qi, L. et al. (2022) 'On the capacity of Sentinel-1 synthetic aperture radar in detecting floating macroalgae and other floating matters', Remote Sensing of Environment, 280, 113188.",
    "Savastano, S. et al. (2021) 'A First Approach to the Automatic Detection of Marine Litter in SAR Images Using Artificial Intelligence', IGARSS 2021, pp. 8704-8707.",
    "Topouzelis, K. et al. (2020) 'Remote Sensing of Sea Surface Artificial Floating Plastic Targets with Sentinel-2 and Unmanned Aerial Systems (Plastic Litter Project 2019)', Remote Sensing, 12(12), 2013.",
]
for ref in references:
    p = doc.add_paragraph(ref, style="Normal")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.size = Pt(9.5)

doc.core_properties.title = "Mock Literature Review Section - SAR Marine Debris"
doc.core_properties.subject = "Illustrative literature-review flow from data creation to machine learning"
doc.core_properties.author = ""
doc.core_properties.keywords = "SAR; marine debris; literature review; labelled data; machine learning"

doc.save(OUTPUT)
print(OUTPUT)
