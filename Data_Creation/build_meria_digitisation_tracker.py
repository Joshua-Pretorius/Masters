#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SA_DIR = ROOT / "meria_sa_plastic_s1_slc"
GLOBAL_DIR = ROOT / "meria_global_s1_slc"
OUT_PATH = ROOT / "MERIA_digitisation_tracker.docx"


@dataclass
class TrackerRow:
    dataset: str
    observation_id: str
    area: str
    role: str
    date: str
    planet_acquired: str
    sar_planet_delta_h: str
    status: str
    processed_location: str
    processed_location_target: str
    notes: str
    save_location: str


def slug(text: str) -> str:
    return "".join(ch if ch.isalnum() else "_" for ch in text.strip()).strip("_")


def normalize_local_path(path_text: str) -> str:
    match = re.match(r"^/mnt/([a-zA-Z])/(.*)$", path_text)
    if not match:
        return path_text
    drive = match.group(1).upper()
    tail = match.group(2).replace("/", "\\")
    return f"{drive}:\\{tail}"


def load_rows(csv_path: Path) -> dict[tuple[str, str], dict[str, str]]:
    rows: dict[tuple[str, str], dict[str, str]] = {}
    with csv_path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            rows[(row["obs_id"], row["date"])] = row
    return rows


def load_manifest_lookup(root_dir: Path) -> dict[tuple[str, str], dict[str, str | Path]]:
    processed_root = root_dir / "processed_slc"
    manifests = sorted(processed_root.rglob("*_slc_manifest.json"))
    lookup: dict[tuple[str, str], dict[str, str | Path]] = {}
    for manifest_path in manifests:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8-sig"))
        lookup[(manifest["observation_id"], manifest["role"])] = {
            "status": manifest.get("status", "unknown"),
            "scene_dir": manifest_path.parent,
            "outputs": manifest.get("outputs") or {},
        }
    return lookup


def preferred_processed_path(manifest_info: dict[str, str | Path] | None, scene_dir: Path) -> Path:
    if manifest_info:
        outputs = manifest_info.get("outputs")
        if isinstance(outputs, dict):
            vv_path = outputs.get("vv")
            if vv_path:
                return Path(normalize_local_path(str(vv_path)))
            for output_path in outputs.values():
                if output_path:
                    return Path(normalize_local_path(str(output_path)))
    return scene_dir


def add_hyperlink(paragraph, text: str, target: str, font_size: Pt) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    rel_id = paragraph.part.relate_to(Path(target).as_uri(), RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_props.append(style)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(font_size.pt * 2)))
    run_props.append(size)

    ascii_font = OxmlElement("w:rFonts")
    ascii_font.set(qn("w:ascii"), "Arial")
    ascii_font.set(qn("w:hAnsi"), "Arial")
    run_props.append(ascii_font)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_props)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def tracker_rows(dataset: str, root_dir: Path, csv_rows: dict[tuple[str, str], dict[str, str]]) -> list[TrackerRow]:
    manifest_lookup = load_manifest_lookup(root_dir)
    rows: list[TrackerRow] = []
    for (obs_id, obs_date), row in csv_rows.items():
        for role in ("before", "after"):
            granule_safe = (row.get(f"{role}_name") or "").strip()
            if not granule_safe or granule_safe == "-":
                continue
            granule = granule_safe.removesuffix(".SAFE")
            acq_key = granule.split("_")[5]
            scene_dir = root_dir / "processed_slc" / f"{obs_id}_{slug(row['area'])}" / f"{role}_{acq_key}"
            manifest_info = manifest_lookup.get((obs_id, role))
            status = "not started"
            if manifest_info:
                status = str(manifest_info["status"])
                scene_dir = Path(manifest_info["scene_dir"])
            processed_path = preferred_processed_path(manifest_info, scene_dir)
            save_dir = scene_dir / "digitised_patches"
            save_dir.mkdir(parents=True, exist_ok=True)
            rows.append(
                TrackerRow(
                    dataset=dataset,
                    observation_id=obs_id,
                    area=row["area"],
                    role=role,
                    date=obs_date,
                    planet_acquired=(row.get("planet_acquired") or "-").strip() or "-",
                    sar_planet_delta_h=(row.get(f"{role}_delta_h") or "-").strip() or "-",
                    status=status,
                    processed_location=str(processed_path),
                    processed_location_target=str(processed_path),
                    notes=(row.get("notes") or "").strip(),
                    save_location=str(save_dir),
                )
            )
    return rows


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def build_docx(rows: list[TrackerRow], out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = title.add_run("MERIA Digitisation Tracker")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph(
        "South Africa and global SLC scenes tracked for manual digitisation. "
        "Use the status column to see what is already processed, and fill in the number of patches digitized after review."
    )
    subtitle.style = doc.styles["Normal"]

    table = doc.add_table(rows=1, cols=12)
    table.style = "Table Grid"
    table.autofit = False

    headers = [
        "Dataset",
        "Observation ID",
        "Area",
        "Role",
        "Date",
        "Planet acquired (UTC)",
        "Delta t h (SAR-Planet)",
        "Processed or not",
        "Processed file location",
        "Observation notes",
        "Patches digitized afterwards",
        "Save digitised outputs to",
    ]
    widths = [0.7, 1.45, 1.0, 0.55, 0.82, 1.55, 1.0, 0.95, 1.85, 2.45, 1.0, 1.85]

    for idx, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "D9EAF7")
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9)

    for row in rows:
        cells = table.add_row().cells
        values = [
            row.dataset,
            row.observation_id,
            row.area,
            row.role,
            row.date,
            row.planet_acquired,
            row.sar_planet_delta_h,
            row.status,
            row.processed_location,
            row.notes or "-",
            "",
            row.save_location,
        ]
        for idx, (cell, value, width) in enumerate(zip(cells, values, widths)):
            cell.width = Inches(width)
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8 if idx in {5, 8, 9, 11} else 9)
            if idx == 8 and row.processed_location_target:
                cell.text = ""
                add_hyperlink(cell.paragraphs[0], row.processed_location, row.processed_location_target, Pt(8))

    doc.save(out_path)


def main() -> None:
    sa_rows = load_rows(SA_DIR / "MERIA_SA_plastic_nearest_S1_SLC_before_after.csv")
    global_rows = load_rows(GLOBAL_DIR / "MERIA_global_plastic_nearest_S1_SLC_before_after.csv")
    rows = tracker_rows("SA", SA_DIR, sa_rows) + tracker_rows("Global", GLOBAL_DIR, global_rows)
    role_order = {"before": 0, "after": 1}
    rows.sort(key=lambda row: (row.dataset, row.observation_id, role_order.get(row.role, 99), row.date, row.area))
    build_docx(rows, OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
